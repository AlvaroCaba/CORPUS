import PyPDF2
import re
import os
import pandas as pd
from collections import Counter, defaultdict
from ollama import generate
from tqdm import tqdm
from pdf2image import convert_from_path
import cv2
from sklearn.cluster import KMeans, MiniBatchKMeans
import matplotlib.pyplot as plt
import numpy as np
import fitz  # PyMuPDF para extraer imágenes del PDF
from PIL import Image
import io
import torch 
import gc
import csv
from transformers import Blip2Processor, Blip2ForConditionalGeneration, InstructBlipProcessor, InstructBlipForConditionalGeneration, AutoProcessor
import yaml
import re, json
from tqdm import tqdm
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import ast
from docx import Document
from pathlib import Path
import colorsys



# Load prompts from prompts.yml at module level
with open(os.path.join(os.path.dirname(__file__), 'prompts.yml'), 'r', encoding='utf-8') as f:
    data = yaml.safe_load(f)    # <-- carga todo el YAML
    PROMPTS_GET_ADJ = data['prompts_get_adjectives']
    PROMPTS_CHANGE_ADJ = data['prompts_change_adjectives']


def extract_pdf_text(file_path):
    """
    Extracts text from a PDF file.
    
    Parameters:
      file_path (str): Path to the PDF file.
    
    Returns:
      str: Text extracted from all pages of the PDF.
    """
    with open(file_path, "rb") as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


def clean_text(text):
    """
    Cleans the extracted text by:
      1. Joining words split by a hyphen at the end of a line.
      2. Replacing any remaining newline characters with a space.
    
    Parameters:
      text (str): Original text.
    
    Returns:
      str: Cleaned text.
    """
    # 1. Join words that are split with a hyphen at the end of a line:
    text = re.sub(r'(\w+)\s*-\s*\n\s*(\w+)', r'\1\2', text)
    # 2. Replace the remaining newline characters with a space:
    text = re.sub(r'\n+', ' ', text)
    return text


def save_sentences_to_file(text, output_path):
    """
    Splits the input text into individual sentences and writes each sentence as a separate line to the specified output file.
    Parameters:
        text (str): The input string containing text to be split into sentences.
        output_path (str): The file path where the sentences will be saved. The function ensures that the directory exists.
    Returns:
        List[str]: A list of cleaned sentences obtained from the input text, with leading and trailing whitespace removed.
    Notes:
        - The function splits the text into sentences by using a regular expression that detects common sentence-ending punctuation marks.
        - Each non-empty sentence is saved with one sentence per line in the output file in UTF-8 encoding.
    """
    
    # Ensure the directory exists (optional)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Step 1: Split the text into sentences using regex.
    sentences = re.split(r'(?<=[.!?])\s+', text)
    clean_sentences = [sentence.strip() for sentence in sentences if sentence.strip() != '']
    
    # Save the sentences to the specified output file.
    with open(output_path, "w", encoding="utf-8") as file:
        for sentence in clean_sentences:
            file.write(sentence + "\n")
    return clean_sentences


def load_adjective_dictionary(file_path):
    """
    Loads an adjective dictionary from a file.
    
    The file should have one entry per line, where the first token is the variant 
    and the second token is the base form.
    
    Parameters:
      file_path (str): Path to the adjective dictionary file.
      
    Returns:
      dict: A dictionary mapping adjective variants to their base forms.
    """
    adj_map = {}
    with open(file_path, encoding='utf-8') as f:
        for line in f:
            if line.strip():
                parts = line.split()
                variant = parts[0].lower()
                base = parts[1].lower()
                adj_map[variant] = base
    return adj_map


def extract_adjectives_from_text(text, adjectives_map):
    """
    Processes the given text by splitting it into sentences, extracting adjectives,
    and returning a DataFrame with the results.
    
    Parameters:
      text (str): The text to process.
      adjectives_map (dict): A dictionary mapping adjective variants to base forms.
      
    Returns:
      pd.DataFrame: A DataFrame with columns 'sentence', 'adjective_variant', and 'adjective_base'.
    """
    
    # Step 1: Process each sentence and extract adjectives.
    data = []
    for sentence in text:
        words = re.findall(r'\w+', sentence.lower())
        variants_in_sentence = []
        bases_in_sentence = []
        
        for word in words:
            if word in adjectives_map:
                if word not in variants_in_sentence:
                    variants_in_sentence.append(word)
                base = adjectives_map[word]
                if base not in bases_in_sentence:
                    bases_in_sentence.append(base)
                
        data.append({
            'sentence': sentence,
            'adjective_variant': ', '.join(variants_in_sentence) if variants_in_sentence else None,
            'adjective_base': ', '.join(bases_in_sentence) if bases_in_sentence else None
        })
    
    # Create and return a DataFrame with the results.
    df_adjectives = pd.DataFrame(data)
    return df_adjectives


def count_adjective_frequencies(text, adjectives_map):
    """
    Processes the given text to count the frequency of adjectives (in their base form)
    and collects the variant forms encountered in the text.
    
    Parameters:
      text (str): The text to process.
      adjectives_map (dict): A dictionary mapping adjective variants to their base forms.
    
    Returns:
      list of tuples: Each tuple contains (base_adjective, variants_string, count, frequency),
                      sorted by count in descending order.
    """
    # Initialize a dictionary to store, for each base adjective, the set of encountered variants.
    variants_by_base = defaultdict(set)
    
    # Preprocess the text: convert to lowercase and extract words.
    words_in_text = re.findall(r'\w+', text.lower())
    
    # Build a list of adjectives (in base form) found in the text, while storing encountered variants.
    adjectives_in_text = []
    for word in words_in_text:
        if word in adjectives_map:
            base = adjectives_map[word]
            adjectives_in_text.append(base)
            variants_by_base[base].add(word)
    
    # Count the frequency of each base adjective found.
    count = Counter(adjectives_in_text)
    total_adjectives = sum(count.values())
    
    results = []
    # Sort the adjectives by count in descending order and prepare results.
    for base_adjective, cnt in sorted(count.items(), key=lambda x: x[1], reverse=True):
        frequency = cnt / total_adjectives  # frequency as a decimal
        variants = ', '.join(sorted(variants_by_base[base_adjective]))
        results.append((base_adjective, variants, cnt, frequency))
    
    return results, count, variants_by_base


def call_ollama_expression(prompt: str, model: str) -> str:   
    """
    Sends a prompt to the Ollama API and returns the generated response.
    
    Parameters:
        prompt (str): The prompt to send to the API.
        model (str): The name of the model to be used.
    
    Returns:
        str: The generated response from the API.
    """
    try:
        response = generate(
            model=model,
            prompt=prompt,
            options={'temperature': 0.0},
            stream=True
        )
        full_response = ""
        for chunk in response:
            #print(chunk['response'], end='', flush=True)
            full_response += chunk['response']
        return full_response
    except Exception as e:
        print(f"Error generating response: {e}")
        return None

def extract_adjectives_ia(sentence: str, model: str, language: str) -> str:
    """
    Extracts adjectives from a given sentence by invoking an AI model with a language-specific prompt.
    Parameters:
        sentence (str): The input sentence from which adjectives will be extracted.
        model (str): The identifier or configuration string for the AI model to use.
        language (str): The language code used to select the appropriate prompt template.
    Returns:
        str: A comma-separated string of normalized adjectives extracted from the sentence.
            Returns 'none' if no adjectives are found or if any error occurs during processing.
    Raises:
        ValueError: If there is no defined prompt for the provided language.
    """
  
    prompt_template = PROMPTS_GET_ADJ.get(language)
    if not prompt_template:
        raise ValueError(f"No hay prompt definido para el idioma: {language}")

    prompt = prompt_template.replace('{sentence}', sentence)
    raw = call_ollama_expression(prompt, model)

    # 1) aislamos el bloque {...} más externo
    m = re.search(r'\{.*\}', raw, re.S)
    if not m:
        return []          # respuesta corrupta

    try:
        data = json.loads(m.group())
        adjs = data.get('adjectives', [])
    except json.JSONDecodeError:
        return 'none'

    # b) normalizamos, quitamos duplicados y vacíos
    norm = lambda t: str(t).strip().lower()
    clean = sorted({norm(a) for a in adjs if norm(a)})

    return ', '.join(clean) if clean else 'none'


def adjectives_list_to_base_form(adj: str,
                                adjectives_map: dict) -> str:
    """
    Converts a string of adjectives into their corresponding base forms using a provided mapping.

    The function performs the following steps:
    1. Checks if the input adjective string is empty; if so, returns 'none'.
    2. Splits the input string into individual words using commas or whitespace as delimiters and strips extra spaces.
    3. Normalizes each word by converting it to lowercase.
    4. Uses the provided adjectives_map to retrieve and collect the base forms of the adjectives found in the normalized list.
    5. Returns a single string containing the sorted list of base forms joined by commas. If no adjectives are found in the map, returns 'none'.

    Parameters:
        adj (str): A string containing adjectives (possibly separated by commas and/or whitespace).
        adjectives_map (dict): A dictionary mapping normalized adjectives to their base forms.

    Returns:
        str: A comma-separated string of sorted base forms if any are found, otherwise 'none'.
    """

    if not adj:
        return 'none'

    # 1) Pasar la cadena a lista de palabras (quitando espacios vacíos)
    adj_list = [w.strip() for w in re.split(r'[,\s]+', adj) if w.strip()]

    # 2) Normalizador: minúsculas y sin acentos
    norm = lambda t: t.lower()

    # 3) Conjunto de formas base encontradas
    bases = {adjectives_map[norm(w)]
             for w in adj_list
             if norm(w) in adjectives_map}

    # 4) Volvemos a **cadena** (ordenada) o 'none'
    return ', '.join(sorted(bases)) if bases else 'none'



def process_sentences_with_ollama(sentences: list[str],
                                  model: str,
                                  language: str,
                                  adjectives_map: dict,
                                  max_retries: int
                                  ) -> pd.DataFrame:
    """
    Devuelve un DataFrame con:
      • Sentence
      • Adjectives_IA_variant (lista)
      • Adjectives_IA_base   (string “a, b, c” o 'none')
    """
    df = pd.DataFrame(sentences, columns=['sentence'])

    df['Adjectives_IA_variant'] = df['sentence'].apply(
        lambda s: extract_adjectives_ia(s, model, language)
    )

    df['Adjectives_IA_base_dic'] = df.apply(
        lambda row: adjectives_list_to_base_form(
            row['Adjectives_IA_variant'],
            adjectives_map
        ),
        axis=1
    )
    df['Adjectives_IA_base'] = df['Adjectives_IA_variant'].apply(
        lambda adj: parse_adjectives_ia(adj, model, language,max_retries)
    )
    return df

def parse_adjectives_ia(adj: str,
                        model: str,
                        language: str,
                        max_retries,
                        ) -> str:
    """
    Llama a Ollama hasta max_retries veces o hasta que el número de bases
    devueltas coincida con el número de variantes en `adj`. 
    Devuelve 'none' si se agotan los intentos o falla.
    """
    # 1) Prepara la lista de variantes de entrada
    variants = [v.strip() for v in adj.split(',') if v.strip()]
    n = len(variants)
    if n == 0 or adj.lower().strip() == 'none':
        return 'none'

    # 2) Construye el prompt una sola vez
    prompt_template = PROMPTS_CHANGE_ADJ.get(language)
    if not prompt_template:
        raise ValueError(f"No hay prompt definido para {language}")
    prompt = prompt_template.replace('{adjective}', adj)

    # 3) Intenta hasta max_retries
    for attempt in range(1, max_retries + 1):
        raw = call_ollama_expression(prompt, model).strip()

        # extrae el bloque JSON
        m = re.search(r'\{.*?\}', raw, flags=re.S)
        if not m:
            continue

        try:
            data = json.loads(m.group())
        except json.JSONDecodeError:
            continue

        base_field = data.get('base')
        if not base_field:
            continue

        # 4) Normaliza base_field a lista de strings
        if isinstance(base_field, list):
            bases = [str(x).strip().lower() for x in base_field if str(x).strip()]
        else:
            bases = [b.strip().lower() for b in str(base_field).split(',') if b.strip()]

        # 5) Comprueba si la longitud coincide
        if len(bases) == n:
            return ', '.join(bases)

        # opcional: log de diagnóstico
        print(f"[Intento {attempt}] Devueltas {len(bases)} bases, esperaba {n}. Reintentando...")
    print(f"⚠️ Error: No se pudo obtener la base de '{adj}' tras {max_retries} intentos.")
    return 'none'


def parse_adjectives(text):
    """
    Dado un texto (cadena), devuelve un conjunto con cada adjetivo en minúsculas.
    Se asume que los adjetivos están separados por comas.
    Si el texto es 'None' o vacío, devuelve un conjunto vacío.
    """
    # Quitamos espacios y convertimos a minúsculas
    if not text or text.strip().lower() == 'none':
        return set()
    return set(adj.strip() for adj in text.lower().split(',') if adj.strip())

def _count_items(cell):
    """
    Devuelve cuántos adjetivos hay en la celda:
      • lista  -> len(lista)
      • str    -> nº de tokens no vacíos separados por coma o espacio
      • None / NaN / resto -> 0
    """
    if isinstance(cell, list):
        return len(cell)

    if isinstance(cell, str):
        txt = cell.strip()
        if not txt or txt.lower() == 'none':
            return 0
        # divide por comas y/o espacios y cuenta los no vacíos
        return len([w for w in re.split(r'[,\s]+', txt) if w])

    return 0   # para NaN, None, etc.


def calculate_metrics(df):
    # Asegúrate de usar los nombres reales de las columnas (ojo a mayúsculas)
    total_base        = df['adjective_base']       .apply(_count_items).sum()
    total_ia_variant  = df['Adjectives_IA_variant'].apply(_count_items).sum()
    total_ia_base     = df['Adjectives_IA_base']   .apply(_count_items).sum()

    print("Total adjective_base:", total_base)
    print("Total Adjectives_IA_variant:", total_ia_variant)
    print("Total Adjectives_IA_base:", total_ia_base)

    return {
        "adjective_base": total_base,
        "adjectives_IA_variant": total_ia_variant,
        "Adjectives_IA_base": total_ia_base
    }
     

def extract_sentences_by_adjective(sentences, adjectives_count, variants_by_base):
    """
    Iterates over a list of sentences and, for each base adjective (from adjectives_count),
    searches for any of its variants (from variants_by_base) in the sentence.
    
    Parameters:
        sentences (list of str): List of sentences to process.
        adjectives_count (dict): Dictionary whose keys are base adjectives (and values, por ejemplo, su conteo).
        variants_by_base (dict): Dictionary mapping each base adjective to a collection (set/list) of its variants.
    
    Returns:
        dict: A dictionary mapping each base adjective to a list of tuples (sentence_number, sentence)
              where at least one variant was found.
    """
    sentences_by_adjective = defaultdict(list)
    
    # Iteramos por cada oración, mostrando el progreso con tqdm.
    for i, sentence in enumerate(tqdm(sentences, desc="Processing sentences"), start=1):
        sentence_lower = sentence.lower()
        # Para cada adjetivo base presente en el diccionario de conteo...
        for base in adjectives_count.keys():
            # Revisamos cada variante asociada a ese adjetivo base.
            for variant in variants_by_base[base]:
                # Buscamos la variante como palabra completa en la oración.
                if re.search(r'\b' + re.escape(variant) + r'\b', sentence_lower):
                    sentences_by_adjective[base].append((i, sentence))
                    break  # Si se encontró una variante, no es necesario buscar más para este adjetivo.
    
    return sentences_by_adjective


def get_noun_from_row(row, model):
    """
    Constructs a prompt for a given row with single adjective and calls the Ollama API.
    """
    adjective = row['adjective']
    sentence = row['sentence']
    prompt = (
        f"En la siguiente oración, ¿a qué sustantivo hace referencia el adjetivo '{adjective}'?\n"
        f"Oración: {sentence}\n"
        f"En la salida dime solo el sustantivo, no escriba nada más"
    )
    response = call_ollama_expression(prompt, model)
    return response


def process_nouns(df, model, output_file="adjectives_nouns.xlsx"):  
    """
    Processes a DataFrame to extract nouns for each adjective in the 'adjectives' column
    by calling the Ollama API. Explodes rows so each adjective is processed separately.

    Parameters:
        df (pd.DataFrame): The input DataFrame with columns 'sentence' and 'adjectives'.
        model (str): The model to use for the API calls.
        output_file (str): The filename for saving the resulting DataFrame.

    Returns:
        pd.DataFrame: The exploded DataFrame with an additional column 'noun'.
    """
    # Ensure progress_apply is available
    tqdm.pandas(desc="Processing adjectives")

    # Split and explode adjectives into separate rows
    df = df.copy()
    df['adjectives'] = df['adjectives'].fillna('')
    df['adjectives_list'] = df['adjectives'].apply(
        lambda x: [adj.strip() for adj in x.split(',') if adj.strip()]
    )
    df_exploded = df.explode('adjectives_list').reset_index(drop=True)
    df_exploded.rename(columns={'adjectives_list': 'adjective'}, inplace=True)
    # If no adjective (empty string), skip API call and fill noun as empty

    # Apply extraction
    df_exploded['noun'] = df_exploded.progress_apply(
            lambda row: extract_noun(row, model),
            axis=1
        )
    # Save and return
    df_exploded.to_csv(output_file, index=False, sep=';')
    return df_exploded

def extract_noun(row, model):
    adj = row['adjective']
    if not adj:
        return ''
    return get_noun_from_row({'adjective': adj, 'sentence': row['sentence']}, model)

def convert_pdf_to_images(pdf_path, output_folder, dpi=300, image_format="PNG"):
    """
    Converts a PDF file into images and saves them in the specified output folder.

    Parameters:
        pdf_path (str): The path to the PDF file.
        output_folder (str): The directory where the images will be saved.
        dpi (int): The resolution (dots per inch) for the conversion.
        image_format (str): The image format to save (e.g., "PNG").

    Returns:
        int: The total number of pages (images) converted.
    """
    # Create the output folder if it does not exist.
    os.makedirs(output_folder, exist_ok=True)

    # Convert the PDF to images.
    images = convert_from_path(pdf_path, dpi=dpi)
    
    total_pages = len(images)
    print(f"Detected {total_pages} pages in the PDF. Starting conversion...\n")
    
    # Save the generated images using a progress bar.
    for i, img in tqdm(enumerate(images, start=1), total=total_pages, desc="Converting pages"):
        img_path = os.path.join(output_folder, f"page_{i}.{image_format.lower()}")
        img.save(img_path, image_format)
    
    print(f"\nConversion completed successfully. Images saved in '{output_folder}'")
    return total_pages

def get_image_files(folder, extension=".png"):
    """
    Returns a sorted list of image file names from the specified folder with the given extension.
    
    Parameters:
        folder (str): Directory containing the images.
        extension (str): File extension to filter by (default ".png").
        
    Returns:
        list: Sorted list of image file names.
    """
    return sorted([f for f in os.listdir(folder) if f.endswith(extension)])


def extract_pixels(image):
    """
    Converts an image (in BGR format) to a flattened array of RGB pixels.
    
    Parameters:
        image (np.ndarray): The input image in BGR format.
        
    Returns:
        np.ndarray: A 2D array of pixels in RGB.
    """
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    return image_rgb.reshape(-1, 3)


def plot_colors(color_percentages, figsize=(8, 2), edgecolor='black', linewidth=1.5, show_text=True):
    """
    Plots a horizontal bar chart showing the percentage of each dominant color,
    always displaying text on each bar and using an edge color.
    
    Parameters:
        color_percentages (dict): A dictionary with keys as RGB tuples and values as percentages.
        figsize (tuple): Figure size.
        edgecolor (str): Edge color for the bars.
        linewidth (float): Line width for the bar borders.
        show_text (bool): If True, displays the percentage text on each bar.
    """
    fig, ax = plt.subplots(figsize=figsize)
    bars = ax.barh(
        range(len(color_percentages)),
        list(color_percentages.values()),
        color=[np.array(color) / 255 for color in color_percentages.keys()],
        edgecolor=edgecolor,
        linewidth=linewidth
    )
    
    ax.set_yticks([])
    ax.set_xlabel("Percentage")
    
    if show_text:
        for bar in bars:
            width = bar.get_width()
            ax.text(width + 0.5, bar.get_y() + bar.get_height() / 2, f'{width:.2f}%', va='center', ha='left')
            
    plt.show()


def get_color_percentages(image, num_colors=5):
    """
    Converts an image to RGB, reshapes it into a pixel array, performs MiniBatchKMeans clustering
    to identify dominant colors, and calculates the percentage of each color.
    
    Parameters:
        image (np.ndarray): The input image (in BGR format as read by OpenCV).
        num_colors (int): The number of dominant colors to extract.
    
    Returns:
        dict: A dictionary where keys are dominant colors (as RGB tuples) and values are their percentages.
    """
    pixels = extract_pixels(image)
    # Use MiniBatchKMeans for per-image clustering.
    mb_kmeans = MiniBatchKMeans(n_clusters=num_colors, random_state=42, batch_size=1000)
    mb_kmeans.fit(pixels)
    counts = Counter(mb_kmeans.labels_)
    total_pixels = sum(counts.values())
    color_percentages = {
        tuple(map(int, mb_kmeans.cluster_centers_[i])): (count / total_pixels) * 100
        for i, count in counts.items()
    }
    return color_percentages


def analyze_dominant_colors(image_folder, num_colors=20):
    """
    Analyzes each PNG image in the specified folder to extract its dominant colors using MiniBatchKMeans.
    
    Parameters:
        image_folder (str): Path to the folder containing images.
        num_colors (int): Number of dominant colors to extract per image.
    
    Returns:
        list: A list of tuples (filename, color_percentages) for each processed image.
    """
    image_files = get_image_files(image_folder)
    total_pages = len(image_files)
    results = []
    
    for i, file in tqdm(enumerate(image_files, start=1), total=total_pages, desc="Analyzing colors"):
        image_path = os.path.join(image_folder, file)
        img = cv2.imread(image_path)
        if img is None:
            continue
        
        color_percentages = get_color_percentages(img, num_colors=num_colors)
        results.append((file, color_percentages))
    
    print("\n✅ Dominant color analysis per page completed successfully.")
    return results


def train_minibatch_kmeans_on_images(folder, image_files, num_colors=2):
    """
    Trains a MiniBatchKMeans model incrementally on images in the specified folder.
    
    Parameters:
        folder (str): Directory containing the images.
        image_files (list): List of image file names to process.
        num_colors (int): Number of clusters (dominant colors) to extract.
    
    Returns:
        MiniBatchKMeans: The trained MiniBatchKMeans model.
    """
    mb_kmeans = MiniBatchKMeans(n_clusters=num_colors, random_state=42, batch_size=1000)
    
    for file in tqdm(image_files, desc="Adjusting MiniBatchKMeans"):
        image_path = os.path.join(folder, file)
        img = cv2.imread(image_path)
        if img is None:
            print(f"Error loading image {file}")
            continue
        pixels = extract_pixels(img)
        mb_kmeans.partial_fit(pixels)
    
    return mb_kmeans


def calculate_overall_color_percentages(folder, image_files, model):
    """
    Igual que antes, pero mantiene los centroides en float (4 dec.) para
    evitar colisiones entre colores casi idénticos.
    """
    counts = Counter()
    total_pixels = 0

    for file in tqdm(image_files, desc="Calculating overall percentages"):
        img = cv2.imread(os.path.join(folder, file))
        if img is None:
            continue
        pixels = extract_pixels(img)
        labels = model.predict(pixels)
        counts.update(labels)
        total_pixels += len(labels)

    overall_color_percentages = {
        # centroides a 4 decimales → clave única
        tuple(np.round(model.cluster_centers_[i], 4)): (cnt / total_pixels) * 100
        for i, cnt in counts.items()
    }

    # orden descendente
    return dict(sorted(overall_color_percentages.items(),
                       key=lambda kv: kv[1], reverse=True))


def extract_images_from_pdf(pdf_path, output_folder):
    """
    Extracts images from a PDF file and saves them in the specified output folder.
    
    This function performs the following steps:
      1. Creates the output folder if it doesn't exist.
      2. Opens the PDF document using PyMuPDF.
      3. Iterates through each page of the PDF with a progress bar.
      4. For each page, retrieves all images (with full details).
      5. Extracts each image, converts it to RGB (if needed), and saves it as a JPEG file.
      6. Stores the PIL Image objects in a list for further processing.
    
    Parameters:
        pdf_path (str): The path to the PDF file.
        output_folder (str): The folder where extracted images will be saved.
    
    Returns:
        list: A list of PIL Image objects extracted from the PDF.
    """
    # Create the output folder if it does not exist.
    os.makedirs(output_folder, exist_ok=True)
    
    # Open the PDF document.
    doc = fitz.open(pdf_path)
    images = []
    
    # Iterate through each page in the PDF with a progress bar.
    for page_index in range(len(doc)):
        page = doc[page_index]
        image_list = page.get_images(full=True)
        #print(f"Page {page_index + 1}: {len(image_list)} image(s) found")
        
        for img_index, img in enumerate(image_list, start=1):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Load the image into a PIL Image object.
            image = Image.open(io.BytesIO(image_bytes))
            # Convert to RGB if the image is not already in that mode.
            if image.mode != "RGB":
                image = image.convert("RGB")
            
            # Construct a filename and save the image to the output folder.
            image_filename = f"page_{page_index + 1}_img_{img_index}.jpg"
            save_path = os.path.join(output_folder, image_filename)
            image.save(save_path)
            
            # Append the PIL Image to the list for further processing.
            images.append(image)
    
    return images


def load_model_blip(model):
    """
    Loads the BLIP-2 processor and model (flan-t5-xxl) and determines the computing device.
    
    This function performs the following steps:
      1. Checks if a CUDA-enabled GPU is available; otherwise, uses the CPU.
      2. Loads the BLIP-2 processor from the "Salesforce/blip2-flan-t5-xxl" pretrained model.
      3. Loads the BLIP-2 model for conditional generation from the same pretrained model with specified torch_dtype and device_map.
      4. Sets the model to evaluation mode.
    
    Returns:
        tuple: A tuple containing:
            - processor: The BLIP-2 processor.
            - model: The BLIP-2 model.
            - device (str): The computing device ("cuda" or "cpu").
    """
    device = "cuda" if torch.cuda.is_available() else "cpu"
    processor = Blip2Processor.from_pretrained(model)
    model = Blip2ForConditionalGeneration.from_pretrained(
        pretrained_model_name_or_path = model,
        torch_dtype=torch.float16,
        device_map="auto"
    )
    model.eval()  # Set the model to evaluation mode.
    return processor, model, device

def load_model_instruct(model):
    """
    Loads the BLIP-2 processor and model (flan-t5-xxl) and determines the computing device.
    
    This function performs the following steps:
      1. Checks if a CUDA-enabled GPU is available; otherwise, uses the CPU.
      2. Loads the BLIP-2 processor from the "instruct/blip2-flan-t5-xxl" pretrained model.
      3. Loads the BLIP-2 model for conditional generation from the same pretrained model with specified torch_dtype and device_map.
      4. Sets the model to evaluation mode.
    
    Returns:
        tuple: A tuple containing:
            - processor: The BLIP-2 processor.
            - model: The BLIP-2 model.
            - device (str): The computing device ("cuda" or "cpu").
    """
    device = "cuda"
    processor = InstructBlipProcessor.from_pretrained(model)
    model = InstructBlipForConditionalGeneration.from_pretrained(
        pretrained_model_name_or_path = model,
        torch_dtype=torch.float16,
        device_map="auto"
    )
    model.eval()  # Set the model to evaluation mode.
    return processor, model, device

def process_image(image_path, processor, model, device):
    """
    Procesa UNA imagen, genera descripción y libera caché.
    """
    try:
        image = Image.open(image_path).convert("RGB")
    except Exception as e:
        print(f"[ERROR] {image_path}: {e}")
        return None

    prompt = (
        "Provide a detailed description of everything you see in the image. "
        "Mention the main objects, their position, and any relationships between them. "
        "Include details about the setting or background, colors, shapes, and any noteworthy elements. "
        "If there are people, describe their characteristics and actions. "
        "Keep the description clear, thorough, and well-structured."
    )

    inputs = processor(images=image, text=prompt, return_tensors="pt")
    # Mover tensores al dispositivo y a FP16
    for k, v in inputs.items():
        inputs[k] = v.to(device)
    if "pixel_values" in inputs:
        inputs["pixel_values"] = inputs["pixel_values"].to(torch.float16)

    # Generación más ligera: menos beams y longitud moderada
    with torch.no_grad():
        generated_ids = model.generate(
            **inputs,
            max_length=200,          # con 100 tokens suele ser más que suficiente
            min_length=15,
            num_beams=2,             # reducir de 5 a 2 beams
            repetition_penalty=1.2,
            do_sample=False
        )
    text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0].strip()

    # Limpieza
    del inputs, generated_ids
    torch.cuda.empty_cache()
    gc.collect()

    return text


def load_processed_images(csv_filename):
    """
    Load processed image identifiers from a CSV file.

    This function reads a CSV file specified by `csv_filename` and extracts the first column value
    from each row (excluding the header) into a set. Each value is assumed to correspond to an image
    identifier that has already been processed. If the file does not exist, the function returns an empty set.

    Parameters:
        csv_filename (str): The path to a CSV file which contains processed image data. The CSV file should have
                            a header row followed by rows where the first column is the image identifier.

    Returns:
        set: A set of strings representing the processed image identifiers.
    """
    processed = set()
    if os.path.exists(csv_filename):
        with open(csv_filename, newline="", encoding="utf-8") as f:
            reader = csv.reader(f)
            next(reader, None)
            for row in reader:
                if row:
                    processed.add(row[0])
    return processed

def process_images_in_folder(folder_path, processor, model, device, csv_filename):
    """
    Process images in a specified folder using a processor and model, and record the results.

    This function iterates over all image files (with extensions .png, .jpg, .jpeg, .bmp, or .gif)
    in the provided folder. It creates the directory for the CSV file if it does not exist and loads
    a list of already processed images from the CSV file specified by "csv_filename". Each image that
    has not been processed yet is then processed using the given processor, model, and device. If the
    processing produces a valid description, the filename and description pair is added to the result list.

    Parameters:
        folder_path (str): The path to the folder containing image files.
        processor (object): An image processor used to prepare the image for the model.
        model (object): The model used for processing images.
        device (str or torch.device): The computational device (e.g., CPU or GPU) to use for processing.
        csv_filename (str): The file path to a CSV file that tracks processed images. Necessary directories
                            will be created if they don't exist.

    Returns:
        list: A list of lists, each containing:
              - filename (str): The image file name.
              - desc (str): The description generated from processing the image.
    """
    os.makedirs(os.path.dirname(csv_filename), exist_ok=True)
    processed = load_processed_images(csv_filename)
    results = []

    for filename in tqdm(os.listdir(folder_path), desc="Procesando imágenes"):
        if not filename.lower().endswith(('.png','.jpg','jpeg','bmp','gif')):
            continue
        if filename in processed:
            continue

        img_path = os.path.join(folder_path, filename)
        desc = process_image(img_path, processor, model, device)
        if desc:
            results.append([filename, desc])
    return results

def load_model_fp16_offload(model_name: str):
    """
    Carga el InstructBLIP en FP16 con reparto automático de capas entre GPU y CPU.
    Esto duplica prácticamente la memoria disponible para el modelo sin cuantizar los pesos.
    """
    processor = AutoProcessor.from_pretrained(model_name)
    model = InstructBlipForConditionalGeneration.from_pretrained(
        model_name,
        torch_dtype=torch.float16,   # medio byte por valor
        device_map="auto",          # 🤗 mueve capas entre GPU/CPU
        low_cpu_mem_usage=True      # reduce picos de memoria al cargar
    )
    return processor, model, model.device

def save_results_to_csv(results, csv_filename):
    """
    Saves the results to a CSV file.
    
    Parameters:
        results (list): A list of [image_filename, description] pairs.
        csv_filename (str): The path to the output CSV file.
        
    Returns:
        None
    """
    # Abrir el CSV en modo append para no sobreescribir resultados previos
    with open(csv_filename, mode='a', encoding='utf-8', newline='') as csvfile:
        writer = csv.writer(csvfile)
        # Si el archivo está vacío, escribimos el header
        if os.stat(csv_filename).st_size == 0:
            writer.writerow(["Image", "Description"])
        writer.writerows(results)



def extract_final_response(text):
    """
    Removes any content enclosed between <think> and </think> tags from the generated text.
    
    This function uses a regular expression with the DOTALL flag to remove all text 
    between the <think> and </think> tags, then strips any leading or trailing whitespace.
    
    Parameters:
        text (str): The generated text that may contain <think>...</think> sections.
    
    Returns:
        str: The cleaned text with the <think> sections removed.
    """
    result = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
    return result

def process_description(description, model):
    """
    Constructs a prompt using the provided description, calls the API, and processes the response.
    
    The prompt instructs the API to list the objects, people, etc. that appear in the given sentence.
    The API should return a single line with each element in singular form, separated by commas.
    After receiving the response, this function removes any content enclosed in <think> tags.
    
    Parameters:
        description (str): The sentence description to process.
    
    Returns:
        str: The final processed response from the API with <think> content removed,
             or an empty string if the API call fails.
    """
    prompt = (
        f"Tell me the objects, people, etc. that appear in this sentence. "
        "List them in the singular form, output only a single line with each element separated by a comma, "
        "and include only concrete nouns (no verbs, no abstract nouns or concepts). "
        f"The sentence is: {description}"
    )


    response = call_ollama_expression(prompt, model)
    if response:
        response = extract_final_response(response)
    else:
        response = ""
    return response


def count_elements_from_csv(csv_path, column_name):
    """
    Lee el CSV y devuelve:
      - counts: dict de elemento (minúsculas) a su frecuencia
      - order: lista de elementos en orden de primera aparición
    """
    df = pd.read_csv(csv_path)
    counts = {}
    order = []

    for row in df[column_name].dropna():
        for elem in (x.strip() for x in row.split(',')):
            key = elem.lower()
            if key not in counts:
                counts[key] = 0
                order.append(key)
            counts[key] += 1

    return counts, order

def process_folder_elem(input_folder, output_folder, column_name='elementos'):
    """
    Processes CSV files in the given input folder by counting the occurrences of elements found
    in a specified CSV column, and then creates a corresponding Word document for each CSV file
    with the element frequencies.

    Parameters:
        input_folder (str or pathlib.Path): Path to the directory containing CSV files.
        output_folder (str or pathlib.Path): Path to the directory where the output .docx files will be saved.
        column_name (str, optional): Name of the CSV column that contains the elements to be counted.
                                     Defaults to 'elementos'.

    The function:
        - Converts input_folder and output_folder to pathlib.Path objects.
        - Ensures that the output folder exists by creating it if necessary.
        - Iterates over each CSV file in the input folder.
        - Counts the frequency of elements in the specified column using an external helper function.
        - Sorts these elements by frequency in descending order.
        - Generates a Word document for each CSV file that lists each element (with its first letter capitalized)
          alongside its frequency.
        - Saves the document in the output folder with the same base name as the CSV but with a .docx extension.
        - Prints a confirmation message for each saved document.

    Returns:
        None
    """
    input_folder  = Path(input_folder)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    for csv_path in input_folder.glob('*.csv'):
        counts, order = count_elements_from_csv(csv_path, column_name)
        # ordenar por frecuencia descendente
        sorted_items = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)

        # crear documento
        doc = Document()
        doc.add_heading(f'Conteo de elementos: {csv_path.name}', level=0)
        for key, freq in sorted_items:
            # Capitaliza solo la primera letra, el resto en minúscula
            doc.add_paragraph(f"{key.capitalize()}: {freq}")

        salida = output_folder / csv_path.with_suffix('.docx').name
        doc.save(salida)
        print(f'Guardado: {salida}')

def print_element_counts(counts, order):
    """
    Prints the element counts in the preserved order, with the first letter capitalized.

    Parameters:
        counts (dict): A dictionary mapping each element to its frequency.
        order (list): A list of elements in the order they first appeared.
    """
    for key in order:
        print(f"{key.capitalize()}: {counts[key]}")


def extract_adjectives_consensus(sentence, adjectives_base, adjectives_ia_base, model):
    """
    Creates the prompt by combining the original sentence with the adjectives obtained
    from two methods, then calls the API function to verify and return the final adjectives.
    
    Parameters:
        sentence (str): The sentence to analyze.
        adjectives_base (str): A comma-separated string with adjectives obtained by the selection method.
        adjectives_ia_base (str): A comma-separated string with adjectives obtained by the IA process.
        model (str): The name of the model to use in the API.
    
    Returns:
        str: The IA response with the verified adjectives.
    """
    prompt = (
    f"Tengo la siguiente frase: {sentence}. He obtenido dos conjuntos de adjetivos: por un lado, mediante un método de selección con lista: {adjectives_base}; por otro, mediante procesamiento con IA: {adjectives_ia_base}. "
    "Revisa ambos resultados y, basándote en ellos, responde únicamente con una lista de adjetivos presentes en la frase, separados por comas, sin ningún comentario ni explicación adicional. "
    "Si no hay adjetivos, responde 'None'. Ejemplo de salida: guapo, bonito, alto, feo."
    )
    return call_ollama_expression(prompt, model)

def parse_color(raw_color):
    """
    Convierte raw_color (str, tuple o lista) a tupla de enteros (r, g, b).
    """
    if isinstance(raw_color, (tuple, list)):
        return raw_color
    if isinstance(raw_color, str):
        try:
            return ast.literal_eval(raw_color)
        except Exception:
            # Fallback: eliminar paréntesis y dividir
            vals = raw_color.strip("() ").split(",")
            return tuple(int(v) for v in vals)
    raise TypeError(f"Tipo de color no soportado: {type(raw_color)}")


def generate_pdf(results, pdf_path, title_prefix="Page",
                 page_size=A4, margin=50, box_size=50, font_family="Helvetica"):
    """
    Genera un PDF a partir de los resultados de colores dominantes.

    results: lista de tuplas (image_name, {color: porcentaje})
    pdf_path: ruta de salida del PDF
    """
    c = canvas.Canvas(pdf_path, pagesize=page_size)
    width, height = page_size
    line_height = box_size + 15
    x = margin
    y = height - margin

    for image_name, colors in results:
        # Título de sección
        c.setFont(f"{font_family}-Bold", 14)
        c.drawString(x, y, f"{title_prefix}: {image_name}")
        y -= line_height

        for raw_color, pct in colors.items():
            r, g, b = parse_color(raw_color)

            # Dibuja recuadro de color
            c.setFillColorRGB(r/255, g/255, b/255)
            c.rect(x, y - box_size, box_size, box_size, fill=1, stroke=0)

            # Texto del porcentaje
            c.setFillColorRGB(0, 0, 0)
            c.setFont(font_family, 10)
            c.drawString(x + box_size + 10,
                         y - box_size/2 - 5,
                         f"{pct:.2f}%")

            y -= line_height
            if y < margin + box_size:
                c.showPage()
                y = height - margin

        # Espacio extra entre imágenes
        y -= line_height / 2
        if y < margin + box_size:
            c.showPage()
            y = height - margin

    c.save()
    print(f"PDF guardado en: {pdf_path}")


def process_folder(input_folder, output_folder, num_colors=10):
    """
    Recorre cada subcarpeta en input_folder, analiza colores y genera PDF por subcarpeta.
    """
    os.makedirs(output_folder, exist_ok=True)

    for subfolder in tqdm(os.listdir(input_folder), desc="Procesando carpetas"):
        subfolder_path = os.path.join(input_folder, subfolder)
        if not os.path.isdir(subfolder_path):
            continue

        # Analiza colores dominantes
        results = analyze_dominant_colors(subfolder_path, num_colors=num_colors)
        results_sorted = sorted(results, key=lambda item: natural_key(item[0]))

        # Ruta de salida PDF
        pdf_name = f"{subfolder}_dominant_colors.pdf"
        pdf_path = os.path.join(output_folder, pdf_name)

        # Genera PDF
        generate_pdf(results_sorted, pdf_path)

def natural_key(s):
    """Llave de ordenación natural que separa dígitos y texto."""
    parts = re.split(r"(\d+)", s)
    return tuple(int(p) if p.isdigit() else p.lower() for p in parts)


def complete_analysis(root_input_folder, root_output_folder, num_colors=10):
    """
    Recorre cada carpeta dentro de root_input_folder (cada PDF convertido a imágenes), realiza análisis completo
    y guarda un PDF resumen de colores globales por carpeta.
    """
    os.makedirs(root_output_folder, exist_ok=True)

    for subfolder in tqdm(os.listdir(root_input_folder), desc="Procesando PDFs completos"):
        subfolder_path = os.path.join(root_input_folder, subfolder)
        if not os.path.isdir(subfolder_path):
            continue

        # Obtiene las imágenes de esta carpeta
        image_files = get_image_files(subfolder_path)
        total_pages = len(image_files)
        print(f"\n{subfolder}: Detectadas {total_pages} páginas para extracción de color dominante.")

        # Entrena modelo KMeans con todas las imágenes de la carpeta
        model = train_minibatch_kmeans_on_images(subfolder_path,
                                                  image_files,
                                                  num_colors=num_colors)

        # Carpeta de salida específica para este PDF
        output_subfolder = os.path.join(root_output_folder, subfolder)
        os.makedirs(output_subfolder, exist_ok=True)

        # ➊ Calcula colores globales
        overall_colors = calculate_overall_color_percentages(subfolder_path,
                                                            image_files,
                                                            model)

        # ➋ Clasificación cálido / frío
        warm_cool = color_classification(overall_colors)

        # ➌ Genera PDF e INCLUYE el resumen warm_cool
        generate_overall_pdf(
            overall_colors,
            pdf_path=os.path.join(output_subfolder, "overall_dominant_colors.pdf"),
            title=f"Dominant Colors – {subfolder}",
            warm_cool=warm_cool          # ← nuevo parámetro
        )

def generate_overall_pdf(color_percentages, pdf_path,
                         title="Overall Dominant Colors",
                         warm_cool=None,               # ← nuevo parámetro
                         page_size=A4, margin=50,
                         box_size=50, font_family="Helvetica"):
    """
    Genera un PDF mostrando los colores dominantes con sus porcentajes y,
    opcionalmente, un resumen cálido/frío.
    """
    c = canvas.Canvas(pdf_path, pagesize=page_size)
    width, height = page_size
    line_height = box_size + 15
    x = margin
    y = height - margin

    # ─── Título ─────────────────────────────────────────────────────────
    c.setFont(f"{font_family}-Bold", 16)
    c.drawString(x, y, title)
    y -= line_height * 1.5

    # ─── Lista de colores ───────────────────────────────────────────────
    for raw_color, pct in color_percentages.items():
        r, g, b = parse_color(raw_color)
        # rectángulo
        c.setFillColorRGB(r/255, g/255, b/255)
        c.rect(x, y - box_size, box_size, box_size, fill=1, stroke=0)
        # porcentaje
        c.setFillColorRGB(0, 0, 0)
        c.setFont(font_family, 12)
        c.drawString(x + box_size + 10,
                     y - box_size/2 - 6,
                     f"{pct:.2f}%")
        y -= line_height
        if y < margin + (2 * line_height):        # deja hueco para el resumen
            c.showPage()
            y = height - margin

    # ─── Resumen cálido / frío ──────────────────────────────────────────
    if warm_cool is not None:
        if y < margin + (2 * line_height):        # salto de página si hace falta
            c.showPage()
            y = height - margin

        c.setFont(f"{font_family}-Bold", 14)
        c.drawString(x, y, "Warm vs. Cool Summary")
        y -= line_height * 1.2

        c.setFont(font_family, 12)
        c.drawString(x, y, f"Warm colors: {warm_cool['warm']:.2f}%")
        y -= line_height * 0.8
        c.drawString(x, y, f"Cool colors: {warm_cool['cool']:.2f}%")

    c.save()
    print(f"PDF generado: {pdf_path}")

def color_classification(color_percentages,
                         warm_hue_ranges=((0, 60), (300, 360))):
    """
    Devuelve {'warm': …, 'cool': …} a partir de color_percentages.
    Acepta centroides float o int.
    """
    warm_pct, cool_pct = 0.0, 0.0

    for (r, g, b), pct in color_percentages.items():
        # convierte cada componente a 0-1
        h, s, v = colorsys.rgb_to_hsv(r / 255.0, g / 255.0, b / 255.0)
        hue_deg = h * 360
        is_warm = any(lo <= hue_deg < hi for lo, hi in warm_hue_ranges)

        if is_warm:
            warm_pct += pct
        else:
            cool_pct += pct

    # redondeo final (la suma ≈ 100 %)
    return {"warm": round(warm_pct, 2),
            "cool": round(cool_pct, 2)}